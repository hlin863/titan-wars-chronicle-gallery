from docx import Document
from docx.shared import Pt

from app import update_manuscript_chapter


def test_chapter_save_preserves_font_and_inline_emphasis(tmp_path):
    manuscript = tmp_path / "Manuscript.docx"
    document = Document()
    document.add_heading("Part I", level=1)
    document.add_heading("Formatting Test", level=2)

    paragraph = document.add_paragraph()
    first = paragraph.add_run("Normal ")
    first.font.name = "Garamond"
    first.font.size = Pt(11)

    emphasis = paragraph.add_run("italic")
    emphasis.font.name = "Garamond"
    emphasis.font.size = Pt(11)
    emphasis.italic = True

    ending = paragraph.add_run(" ending")
    ending.font.name = "Garamond"
    ending.font.size = Pt(11)

    document.save(manuscript)

    update_manuscript_chapter(
        manuscript,
        "chapter-1-Formatting-Test",
        ["Normal revised italic ending"],
    )

    reopened = Document(manuscript)
    saved = reopened.paragraphs[2]

    assert saved.text == "Normal revised italic ending"
    assert saved.runs[0].text == "Normal revised "
    assert saved.runs[0].font.name == "Garamond"
    assert saved.runs[0].font.size.pt == 11

    italic_run = next(run for run in saved.runs if run.text == "italic")
    assert italic_run.font.name == "Garamond"
    assert italic_run.font.size.pt == 11
    assert italic_run.italic is True
