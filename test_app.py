import base64
import json
import os
import threading
from pathlib import Path
from urllib.error import HTTPError
from urllib.request import Request, urlopen

from docx import Document

from app import (
    GalleryImage,
    GalleryServer,
    ManuscriptChapter,
    extract_gallery,
    group_chapters,
    render_page,
    update_manuscript_chapter,
)


TEST_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9Z6J0AAAAASUVORK5CYII="
)


def test_gallery_renders_lightbox_and_primary_navigation():
    source = Path("Manuscript.docx")
    images = [
        GalleryImage(
            id=f"image-{index}",
            file_name=f"illustration-{index}.png",
            year=1850 + index,
            chapter=f"Chapter {index}",
            part="Part I",
            context="A scene",
            document_order=index,
            alt_text=f"Illustration {index}",
        )
        for index in range(1, 4)
    ]

    assert group_chapters(images)
    page = render_page(source, images, "test-version").decode("utf-8")

    assert "Illustrated chronology" in page
    assert page.count('class="image-frame"') == 3
    assert 'id="lightbox-previous"' in page
    assert 'id="lightbox-next"' in page
    assert 'id="lightbox-position"' in page
    assert "/media/illustration-1.png?v=test-version" in page
    assert 'href="/"' in page
    assert 'href="/prompt-studio"' in page
    assert 'aria-current="page"' in page


def test_prompt_studio_template_has_return_navigation():
    page = Path("templates/prompt_studio.html").read_text(encoding="utf-8")

    assert "Prompt Studio" in page
    assert 'href="/"' in page
    assert 'href="/prompt-studio"' in page
    assert 'aria-label="Primary navigation"' in page
    assert 'aria-current="page"' in page


def test_manuscript_chapter_text_joins_paragraphs():
    chapter = ManuscriptChapter(
        id="chapter-1",
        title="Arrival",
        part="Part I",
        year=1856,
        metadata_lines=["London, 1856"],
        paragraphs=["First.", "Second."],
        document_order=1,
    )

    assert chapter.text == "First.\n\nSecond."


def test_update_manuscript_chapter_only_replaces_selected_chapter(tmp_path):
    manuscript = tmp_path / "Manuscript.docx"
    document = Document()
    document.add_heading("Part I", level=1)
    document.add_heading("First", level=2)
    document.add_paragraph("First chapter text.")
    document.add_heading("Second", level=2)
    document.add_paragraph("Second chapter text.")
    document.save(manuscript)

    updated = update_manuscript_chapter(
        manuscript, "chapter-2-Second", ["Edited second chapter text."]
    )

    reopened = Document(manuscript)
    assert updated.paragraphs == ["Edited second chapter text."]
    assert [paragraph.text for paragraph in reopened.paragraphs] == [
        "Part I",
        "First",
        "First chapter text.",
        "Second",
        "Edited second chapter text.",
    ]


def test_extract_gallery_syncs_files_without_rebuilding_directory(tmp_path):
    manuscript = tmp_path / "Manuscript.docx"
    source_image = tmp_path / "source.png"
    source_image.write_bytes(TEST_PNG)

    document = Document()
    document.add_heading("Part I", level=1)
    document.add_heading("The Test Chapter", level=2)
    document.add_paragraph("London, 1856")
    document.add_picture(str(source_image))
    document.save(manuscript)

    media_dir = tmp_path / "gallery_media"
    first_dir, first_images = extract_gallery(manuscript, media_dir)

    assert first_dir == media_dir.resolve()
    assert len(first_images) == 1
    stored_image = media_dir / first_images[0].file_name
    assert stored_image.exists()
    assert stored_image.read_bytes() == TEST_PNG

    fixed_time_ns = 1_600_000_000_000_000_000
    os.utime(stored_image, ns=(fixed_time_ns, fixed_time_ns))
    stale_image = media_dir / "stale-image.png"
    stale_image.write_bytes(b"obsolete")
    unrelated_file = media_dir / "notes.txt"
    unrelated_file.write_text("keep me", encoding="utf-8")

    document.add_paragraph("Updated manuscript text with the same illustration.")
    document.save(manuscript)
    second_dir, second_images = extract_gallery(manuscript, media_dir)

    assert second_dir == media_dir.resolve()
    assert len(second_images) == 1
    assert second_images[0].file_name == first_images[0].file_name
    assert stored_image.stat().st_mtime_ns == fixed_time_ns
    assert stored_image.read_bytes() == TEST_PNG
    assert not stale_image.exists()
    assert unrelated_file.exists()
    assert not (media_dir / "manifest.json").exists()


def test_generate_prompt_builds_non_streaming_ollama_chat(monkeypatch):
    import app

    request = {}

    def fake_request(endpoint, payload):
        request.update(endpoint=endpoint, payload=payload)
        return {
            "model": "gemma3:4b",
            "message": {"content": "A grounded visual prompt."},
            "prompt_eval_count": 42,
            "eval_count": 8,
        }

    monkeypatch.setattr(app, "ollama_request", fake_request)
    result = app.generate_prompt_with_ollama(
        model="gemma3:4b",
        chapter_title="Arrival",
        part="Part I",
        year=1856,
        metadata_lines=["London"],
        selected_text="A sufficiently long selected passage.",
        continuity_notes="Keep the blue coat.",
        style_direction="Candlelight.",
    )

    assert request["endpoint"] == "/api/chat"
    assert request["payload"]["stream"] is False
    assert request["payload"]["model"] == "gemma3:4b"
    assert result["prompt"] == "A grounded visual prompt."


def test_prompt_studio_and_chapter_api_are_served(tmp_path):
    manuscript = tmp_path / "Manuscript.docx"
    document = Document()
    document.add_heading("Part I", level=1)
    document.add_heading("The Test Chapter", level=2)
    document.add_paragraph("London, 1856")
    document.add_paragraph("A chapter paragraph long enough to load in the reader.")
    document.save(manuscript)

    gallery = GalleryServer(manuscript)

    from http.server import ThreadingHTTPServer

    server = ThreadingHTTPServer(("127.0.0.1", 0), gallery.handler_class())
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()

    try:
        base_url = f"http://127.0.0.1:{server.server_port}"

        with urlopen(f"{base_url}/prompt-studio") as response:
            studio_page = response.read().decode("utf-8")
            assert response.status == 200
            assert "Prompt Studio" in studio_page
            assert 'href="/"' in studio_page

        with urlopen(f"{base_url}/api/chapters") as response:
            payload = json.loads(response.read().decode("utf-8"))
            assert response.status == 200
            assert payload["ok"] is True
            assert payload["chapters"][0]["title"] == "The Test Chapter"

        save_request = Request(
            f"{base_url}/api/chapters/save",
            data=json.dumps(
                {
                    "chapter_id": payload["chapters"][0]["id"],
                    "version": payload["version"],
                    "paragraphs": ["London, 1857", "Edited chapter paragraph."],
                }
            ).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        with urlopen(save_request) as response:
            saved = json.loads(response.read().decode("utf-8"))
            assert response.status == 200
            assert saved["chapter"]["paragraphs"][-1] == "Edited chapter paragraph."

        stale_request = Request(
            f"{base_url}/api/chapters/save",
            data=json.dumps(
                {
                    "chapter_id": payload["chapters"][0]["id"],
                    "version": payload["version"],
                    "paragraphs": ["London, 1858", "Stale edit."],
                }
            ).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        try:
            urlopen(stale_request)
            raise AssertionError("A stale save should fail")
        except HTTPError as error:
            assert error.code == 409
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=2)
