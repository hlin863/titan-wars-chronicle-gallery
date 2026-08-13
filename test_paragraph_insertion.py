from docx import Document
from docx.shared import Pt

from paragraph_editing import update_manuscript_chapter_with_inserts


def test_insert_paragraph_between_existing_paragraphs_preserves_order_and_format(tmp_path):
    manuscript = tmp_path / "Manuscript.docx"
    document = Document()
    document.add_heading("Part I", level=1)
    document.add_heading("Insertion Test", level=2)

    first = document.add_paragraph()
    first_run = first.add_run("First paragraph.")
    first_run.font.name = "Garamond"
    first_run.font.size = Pt(11)

    second = document.add_paragraph()
    second_run = second.add_run("Second paragraph.")
    second_run.font.name = "Garamond"
    second_run.font.size = Pt(11)
    document.save(manuscript)

    updated = update_manuscript_chapter_with_inserts(
        manuscript,
        "chapter-1-Insertion-Test",
        ["First paragraph.", "Inserted paragraph.", "Second paragraph."],
    )

    reopened = Document(manuscript)
    body = reopened.paragraphs[2:]
    assert updated.paragraphs == [
        "First paragraph.",
        "Inserted paragraph.",
        "Second paragraph.",
    ]
    assert [paragraph.text for paragraph in body] == updated.paragraphs
    assert body[1].runs[0].font.name == "Garamond"
    assert body[1].runs[0].font.size.pt == 11


def test_inserted_paragraph_cannot_be_blank(tmp_path):
    manuscript = tmp_path / "Manuscript.docx"
    document = Document()
    document.add_heading("Part I", level=1)
    document.add_heading("Insertion Test", level=2)
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    document.save(manuscript)

    try:
        update_manuscript_chapter_with_inserts(
            manuscript,
            "chapter-1-Insertion-Test",
            ["First paragraph.", "", "Second paragraph."],
        )
        raise AssertionError("Blank inserted paragraphs should fail")
    except ValueError as error:
        assert "cannot be empty" in str(error)
