from __future__ import annotations

import argparse
import hashlib
import html
import json
import mimetypes
import os
import re
import shutil
import threading
from dataclasses import asdict, dataclass
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Iterable
from urllib.error import HTTPError, URLError
from urllib.parse import parse_qs, unquote, urlparse
from urllib.request import Request, urlopen

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

BASE_DIR = Path(__file__).resolve().parent
DEFAULT_DOCX = Path(os.environ.get("TITAN_WARS_DOCX", "files/Manuscript.docx"))
CACHE_ROOT = BASE_DIR / "instance" / "gallery_cache"
YEAR_RE = re.compile(r"(?<!\d)(17\d{2}|18\d{2}|19\d{2}|20\d{2})(?!\d)")

OLLAMA_BASE_URL = os.environ.get("OLLAMA_BASE_URL", "http://127.0.0.1:11434").rstrip("/")
OLLAMA_DEFAULT_MODEL = os.environ.get("OLLAMA_MODEL", "gemma3:4b")
MAX_SELECTION_LENGTH = 16_000
MAX_DIRECTION_LENGTH = 3_000


@dataclass
class GalleryImage:
    id: str
    file_name: str
    year: int
    chapter: str
    part: str
    context: str
    document_order: int
    alt_text: str


@dataclass
class ManuscriptChapter:
    id: str
    title: str
    part: str
    year: int
    metadata_lines: list[str]
    paragraphs: list[str]
    document_order: int

    @property
    def text(self) -> str:
        return "\n\n".join(self.paragraphs)


PROMPT_WRITER_SYSTEM = """
You are the visual-development prompt writer for Titan Wars, a historical
alternate-history and Victorian gothic novel.

Transform the selected manuscript passage into one precise, production-ready
image-generation prompt.

Rules:
1. Use only facts supported by the passage, chapter metadata, and user continuity notes.
2. Do not invent relationships, motives, injuries, ages, costumes, locations,
   supernatural forms, or plot events.
3. Preserve named characters and distinguish narrators, memories, portraits,
   dreams, and background figures.
4. Preserve historical year, setting, age, body type, clothing, ethnicity,
   emotional condition, and established continuity.
5. Choose one visually coherent moment rather than combining several scenes.
6. Include subject, appearance, posture/action, clothing, setting, composition,
   lighting, emotional tone, restrained background detail, and negative constraints.
7. Do not explain your reasoning or mention the model.
8. Return only the final image prompt.
""".strip()


def iter_blocks(parent: _Document) -> Iterable[Paragraph | Table]:
    for child in parent.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def compact_text(text: str) -> str:
    text = " ".join(text.replace("\xa0", " ").split())
    for repeats in (3, 2):
        if text and len(text) % repeats == 0:
            piece = text[: len(text) // repeats]
            if piece * repeats == text:
                return piece.strip()
    return text


def safe_slug(value: str) -> str:
    value = re.sub(r"[^a-zA-Z0-9_-]+", "-", value).strip("-")
    return value[:70] or "image"


def paragraph_images(paragraph: Paragraph):
    results = []
    for drawing in paragraph._p.xpath(".//w:drawing"):
        blips = drawing.xpath(".//a:blip")
        if not blips:
            continue
        rel_id = blips[0].get(qn("r:embed"))
        if not rel_id:
            continue
        doc_props = drawing.xpath(".//wp:docPr")
        alt = ""
        if doc_props:
            alt = doc_props[0].get("descr") or doc_props[0].get("title") or ""
        results.append((rel_id, compact_text(alt)))
    return results


def file_digest(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def file_stat_signature(path: Path) -> tuple[int, int]:
    stat = path.stat()
    return stat.st_size, stat.st_mtime_ns


def cache_key(docx_path: Path, digest: str | None = None) -> str:
    digest = digest or file_digest(docx_path)
    raw = f"{docx_path.resolve()}:{digest}".encode()
    return hashlib.sha256(raw).hexdigest()[:16]


def extract_gallery(docx_path: Path, digest: str | None = None) -> tuple[Path, list[GalleryImage]]:
    if not docx_path.exists():
        raise FileNotFoundError(f"Manuscript not found: {docx_path}")

    cache_dir = CACHE_ROOT / cache_key(docx_path, digest)
    media_dir = cache_dir / "media"
    manifest_path = cache_dir / "manifest.json"
    if manifest_path.exists():
        data = json.loads(manifest_path.read_text(encoding="utf-8"))
        return media_dir, [GalleryImage(**item) for item in data]

    if cache_dir.exists():
        shutil.rmtree(cache_dir)
    media_dir.mkdir(parents=True, exist_ok=True)

    doc = Document(docx_path)
    images: list[GalleryImage] = []
    current_part = "The 1850s Manuscript"
    current_chapter = "Opening"
    current_year: int | None = None
    lines_since_heading = 0
    recent_lines: list[str] = []
    order = 0

    for block in iter_blocks(doc):
        if not isinstance(block, Paragraph):
            continue
        text = compact_text(block.text)
        style = (block.style.name or "").lower()
        if text:
            if style.startswith("heading 1"):
                current_part = text
            elif style.startswith("heading 2"):
                current_chapter = text
                current_year = None
                lines_since_heading = 0
            else:
                lines_since_heading += 1
            years = [int(x) for x in YEAR_RE.findall(text)]
            if years and len(text) <= 140 and lines_since_heading <= 10:
                current_year = years[-1]
            recent_lines.append(text)
            recent_lines = recent_lines[-8:]

        for rel_id, alt in paragraph_images(block):
            image_part = doc.part.related_parts[rel_id]
            suffix = Path(str(image_part.partname)).suffix.lower() or ".png"
            if suffix not in {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tif", ".tiff"}:
                suffix = ".png"
            order += 1
            year = current_year or 0
            file_name = f"{order:03d}-{year or 'undated'}-{safe_slug(current_chapter)}{suffix}"
            (media_dir / file_name).write_bytes(image_part.blob)
            context_candidates = [line for line in reversed(recent_lines) if line != current_chapter]
            context = next((line for line in context_candidates if len(line) <= 130), "")
            images.append(
                GalleryImage(
                    id=f"image-{order}",
                    file_name=file_name,
                    year=year,
                    chapter=current_chapter,
                    part=current_part,
                    context=context,
                    document_order=order,
                    alt_text=alt or f"Illustration from {current_chapter}",
                )
            )

    images.sort(key=lambda item: (item.year if item.year else 9999, item.document_order))
    manifest_path.write_text(
        json.dumps([asdict(x) for x in images], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return media_dir, images


def extract_chapters(docx_path: Path) -> list[ManuscriptChapter]:
    if not docx_path.exists():
        raise FileNotFoundError(f"Manuscript not found: {docx_path}")

    document = Document(docx_path)
    chapters: list[ManuscriptChapter] = []
    current_part = "The 1850s Manuscript"
    current_title: str | None = None
    current_year = 0
    current_metadata: list[str] = []
    current_paragraphs: list[str] = []
    document_order = 0
    lines_after_heading = 0

    def finish_chapter() -> None:
        nonlocal document_order
        if not current_title:
            return
        document_order += 1
        chapters.append(
            ManuscriptChapter(
                id=f"chapter-{document_order}-{safe_slug(current_title)}",
                title=current_title,
                part=current_part,
                year=current_year,
                metadata_lines=current_metadata.copy(),
                paragraphs=current_paragraphs.copy(),
                document_order=document_order,
            )
        )

    for block in iter_blocks(document):
        if not isinstance(block, Paragraph):
            continue
        text = compact_text(block.text)
        style = (block.style.name or "").lower()
        if style.startswith("heading 1") and text:
            current_part = text
            continue
        if style.startswith("heading 2") and text:
            finish_chapter()
            current_title = text
            current_year = 0
            current_metadata = []
            current_paragraphs = []
            lines_after_heading = 0
            continue
        if not current_title or not text:
            continue
        lines_after_heading += 1
        years = YEAR_RE.findall(text)
        if lines_after_heading <= 5 and len(text) <= 160:
            current_metadata.append(text)
            if years:
                current_year = int(years[-1])
        current_paragraphs.append(text)

    finish_chapter()
    return chapters


def ollama_request(endpoint: str, payload: dict | None = None, timeout: int = 180) -> dict:
    url = f"{OLLAMA_BASE_URL}{endpoint}"
    body = None
    method = "GET"
    if payload is not None:
        body = json.dumps(payload).encode("utf-8")
        method = "POST"
    request = Request(url, data=body, method=method, headers={"Content-Type": "application/json"})
    try:
        with urlopen(request, timeout=timeout) as response:
            return json.loads(response.read().decode("utf-8"))
    except HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"Ollama returned HTTP {exc.code}: {detail}") from exc
    except URLError as exc:
        raise RuntimeError(
            f"Cannot connect to Ollama at {OLLAMA_BASE_URL}. Start Ollama and try again."
        ) from exc


def list_ollama_models() -> list[dict]:
    result = ollama_request("/api/tags")
    models = []
    for item in result.get("models", []):
        name = item.get("name") or item.get("model")
        if not name:
            continue
        details = item.get("details") or {}
        models.append(
            {
                "name": name,
                "parameter_size": details.get("parameter_size", ""),
                "quantization": details.get("quantization_level", ""),
                "family": details.get("family", ""),
                "size": item.get("size", 0),
            }
        )
    return models


def generate_prompt_with_ollama(
    *,
    model: str,
    chapter_title: str,
    part: str,
    year: int,
    metadata_lines: list[str],
    selected_text: str,
    continuity_notes: str,
    style_direction: str,
) -> dict:
    selected_text = selected_text.strip()
    if len(selected_text) < 20:
        raise ValueError("Select at least 20 characters from the chapter.")
    if len(selected_text) > MAX_SELECTION_LENGTH:
        raise ValueError(f"Selection exceeds the {MAX_SELECTION_LENGTH:,}-character limit.")

    context = {
        "chapter": chapter_title,
        "part": part,
        "year": year or None,
        "chapter_metadata": metadata_lines,
        "selected_passage": selected_text,
        "continuity_instructions": continuity_notes.strip()[:MAX_DIRECTION_LENGTH],
        "style_direction": style_direction.strip()[:MAX_DIRECTION_LENGTH],
    }
    response = ollama_request(
        "/api/chat",
        {
            "model": model,
            "stream": False,
            "messages": [
                {"role": "system", "content": PROMPT_WRITER_SYSTEM},
                {
                    "role": "user",
                    "content": "Create an image prompt from this source material:\n\n"
                    + json.dumps(context, ensure_ascii=False, indent=2),
                },
            ],
            "options": {"temperature": 0.35, "top_p": 0.9, "num_predict": 1400},
            "keep_alive": "10m",
        },
    )
    prompt = response.get("message", {}).get("content", "").strip()
    if not prompt:
        raise RuntimeError("Ollama returned an empty prompt.")
    return {
        "prompt": prompt,
        "model": response.get("model", model),
        "prompt_tokens": response.get("prompt_eval_count"),
        "output_tokens": response.get("eval_count"),
        "duration_ns": response.get("total_duration"),
    }


def group_chapters(images: list[GalleryImage]):
    groups, index = [], {}
    for image in images:
        key = (image.year, image.part, image.chapter)
        if key not in index:
            index[key] = {
                "year": image.year,
                "part": image.part,
                "chapter": image.chapter,
                "context": image.context,
                "images": [],
            }
            groups.append(index[key])
        index[key]["images"].append(image)
    return groups


def e(value) -> str:
    return html.escape(str(value), quote=True)


def render_page(source: Path, images: list[GalleryImage], version: str = "") -> bytes:
    groups = group_chapters(images)
    years = sorted({x.year for x in images if x.year})
    year_buttons = "".join(
        f'<button class="year-pill" data-year="{y}">{y}</button>' for y in years
    )
    chapters = []
    for group in groups:
        frames = []
        for image in group["images"]:
            url = f"/media/{e(image.file_name)}?v={e(version)}"
            frames.append(
                f'''<button class="image-frame" type="button" data-full="{url}" data-alt="{e(image.alt_text)}" data-title="{e(group['chapter'])}" aria-label="Open illustration from {e(group['chapter'])}"><img src="{url}" alt="{e(image.alt_text)}" loading="lazy"><span class="zoom-label">View image</span></button>'''
            )
        search_text = f"{group['chapter']} {group['part']} {group['context']} {group['year']}".lower()
        context = f'<div class="context">{e(group["context"])}</div>' if group["context"] else ""
        chapters.append(
            f'''<article class="chapter" data-year="{group['year']}" data-search="{e(search_text)}"><div class="timeline-marker"><span>{group['year'] or '—'}</span></div><div class="chapter-card"><div class="chapter-heading"><p>{e(group['part'])}</p><h2>{e(group['chapter'])}</h2>{context}</div><div class="image-grid count-{len(frames)}">{''.join(frames)}</div></div></article>'''
        )

    template = (BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")
    replacements = {
        "{{MANUSCRIPT_NAME}}": e(source.stem),
        "{{IMAGE_COUNT}}": str(len(images)),
        "{{CHAPTER_COUNT}}": str(len(groups)),
        "{{YEAR_BUTTONS}}": year_buttons,
        "{{CHAPTERS}}": "".join(chapters),
        "{{GALLERY_VERSION}}": e(version),
    }
    for token, value in replacements.items():
        template = template.replace(token, value)
    return template.encode("utf-8")


class GalleryServer:
    def __init__(self, docx_path: Path):
        self.source = docx_path.resolve()
        self.lock = threading.RLock()
        self.stat_signature: tuple[int, int] | None = None
        self.version = ""
        self.media_dir = CACHE_ROOT
        self.images: list[GalleryImage] = []
        self.chapters: list[ManuscriptChapter] = []
        self.index_html = b""
        self.refresh(force=True)

    def refresh(self, force: bool = False) -> bool:
        with self.lock:
            current_stat = file_stat_signature(self.source)
            if not force and current_stat == self.stat_signature:
                return False
            digest = file_digest(self.source)
            if not force and digest == self.version:
                self.stat_signature = current_stat
                return False

            media_dir, images = extract_gallery(self.source, digest)
            chapters = extract_chapters(self.source)
            index_html = render_page(self.source, images, digest[:16])
            self.media_dir = media_dir
            self.images = images
            self.chapters = chapters
            self.index_html = index_html
            self.version = digest
            self.stat_signature = current_stat
            print(
                f"Gallery refreshed: {len(images)} images, {len(chapters)} chapters "
                f"from {self.source.name} (version {digest[:12]})"
            )
            return True

    def refresh_if_changed(self) -> bool:
        try:
            return self.refresh(force=False)
        except (OSError, PermissionError, ValueError, json.JSONDecodeError) as exc:
            print(f"Manuscript refresh deferred: {exc}")
            return False

    def handler_class(self):
        server_state = self

        class Handler(BaseHTTPRequestHandler):
            def send_content(self, payload: bytes, content_type: str, cache_control: str = "no-cache"):
                self.send_response(200)
                self.send_header("Content-Type", content_type)
                self.send_header("Cache-Control", cache_control)
                self.send_header("Content-Length", str(len(payload)))
                self.end_headers()
                self.wfile.write(payload)

            def send_json(self, payload: dict | list, status: int = 200):
                encoded = json.dumps(payload, ensure_ascii=False).encode("utf-8")
                self.send_response(status)
                self.send_header("Content-Type", "application/json; charset=utf-8")
                self.send_header("Cache-Control", "no-store")
                self.send_header("Content-Length", str(len(encoded)))
                self.end_headers()
                self.wfile.write(encoded)

            def read_json(self) -> dict:
                try:
                    length = int(self.headers.get("Content-Length", "0"))
                except ValueError as exc:
                    raise ValueError("Invalid request length.") from exc
                if length <= 0:
                    raise ValueError("The request body is empty.")
                if length > 100_000:
                    raise ValueError("The request body is too large.")
                raw = self.rfile.read(length)
                try:
                    payload = json.loads(raw.decode("utf-8"))
                except (UnicodeDecodeError, json.JSONDecodeError) as exc:
                    raise ValueError("Invalid JSON request.") from exc
                if not isinstance(payload, dict):
                    raise ValueError("The request body must be a JSON object.")
                return payload

            def send_file(self, root: Path, relative: str, cache_control: str = "no-cache"):
                target = (root / relative).resolve()
                try:
                    target.relative_to(root.resolve())
                except ValueError:
                    self.send_error(403)
                    return
                if not target.is_file():
                    self.send_error(404)
                    return
                mime = mimetypes.guess_type(target.name)[0] or "application/octet-stream"
                self.send_content(target.read_bytes(), mime, cache_control=cache_control)

            def do_GET(self):
                parsed = urlparse(self.path)
                path = unquote(parsed.path)
                query = parse_qs(parsed.query)

                if path in {"/", "/prompt-studio", "/api/version", "/api/chapters", "/health"}:
                    server_state.refresh_if_changed()

                if path == "/":
                    with server_state.lock:
                        self.send_content(server_state.index_html, "text/html; charset=utf-8", "no-store")
                elif path == "/prompt-studio":
                    self.send_file(BASE_DIR / "templates", "prompt_studio.html", "no-store")
                elif path == "/api/version":
                    if query.get("force") == ["1"]:
                        server_state.refresh(force=True)
                    with server_state.lock:
                        self.send_json(
                            {
                                "version": server_state.version[:16],
                                "images": len(server_state.images),
                                "chapters": len(server_state.chapters),
                            }
                        )
                elif path == "/api/chapters":
                    with server_state.lock:
                        chapters = [asdict(chapter) for chapter in server_state.chapters]
                        version = server_state.version[:16]
                    self.send_json({"ok": True, "version": version, "chapters": chapters})
                elif path == "/api/ollama/models":
                    try:
                        self.send_json(
                            {
                                "ok": True,
                                "default_model": OLLAMA_DEFAULT_MODEL,
                                "models": list_ollama_models(),
                            }
                        )
                    except RuntimeError as exc:
                        self.send_json({"ok": False, "error": str(exc), "models": []}, 503)
                elif path == "/health":
                    with server_state.lock:
                        self.send_json(
                            {
                                "status": "ok",
                                "version": server_state.version[:16],
                                "images": len(server_state.images),
                                "chapters": len(server_state.chapters),
                            }
                        )
                elif path.startswith("/media/"):
                    with server_state.lock:
                        media_dir = server_state.media_dir
                    self.send_file(media_dir, path[len("/media/"):], "public, max-age=31536000, immutable")
                elif path.startswith("/static/"):
                    self.send_file(BASE_DIR / "static", path[len("/static/"):], "no-cache")
                else:
                    self.send_error(404)

            def do_POST(self):
                path = unquote(urlparse(self.path).path)
                try:
                    if path != "/api/ollama/write-prompt":
                        self.send_json({"ok": False, "error": "Unknown API route."}, 404)
                        return
                    payload = self.read_json()
                    try:
                        year = int(payload.get("year") or 0)
                    except (TypeError, ValueError):
                        year = 0
                    metadata_lines = payload.get("metadata_lines") or []
                    if not isinstance(metadata_lines, list):
                        metadata_lines = []
                    result = generate_prompt_with_ollama(
                        model=str(payload.get("model") or OLLAMA_DEFAULT_MODEL),
                        chapter_title=str(payload.get("chapter_title") or ""),
                        part=str(payload.get("part") or ""),
                        year=year,
                        metadata_lines=[str(item) for item in metadata_lines if isinstance(item, str)],
                        selected_text=str(payload.get("selected_text") or ""),
                        continuity_notes=str(payload.get("continuity_notes") or ""),
                        style_direction=str(payload.get("style_direction") or ""),
                    )
                    self.send_json({"ok": True, **result})
                except ValueError as exc:
                    self.send_json({"ok": False, "error": str(exc)}, 400)
                except RuntimeError as exc:
                    self.send_json({"ok": False, "error": str(exc)}, 503)
                except Exception as exc:
                    print(f"Prompt Studio error: {exc}")
                    self.send_json(
                        {"ok": False, "error": "Prompt generation failed. Check the server terminal."},
                        500,
                    )

            def log_message(self, fmt, *args):
                print(f"[{self.log_date_time_string()}] {fmt % args}")

        return Handler


def main():
    parser = argparse.ArgumentParser(
        description="Display manuscript illustrations and provide a local Ollama prompt studio."
    )
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX, help="Path to the manuscript DOCX")
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", type=int, default=int(os.environ.get("PORT", "5000")))
    args = parser.parse_args()
    gallery = GalleryServer(args.docx)
    server = ThreadingHTTPServer((args.host, args.port), gallery.handler_class())
    print(f"Titan Wars gallery: http://{args.host}:{args.port}")
    print(f"Prompt Studio: http://{args.host}:{args.port}/prompt-studio")
    print(f"Watching manuscript for changes: {gallery.source}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nStopping gallery.")
    finally:
        server.server_close()


if __name__ == "__main__":
    main()
